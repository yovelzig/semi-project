import json
from pathlib import Path

import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


DATA_DIR = Path("data")
DATA_DIR.mkdir(exist_ok=True)


def create_txt():
    path = DATA_DIR / "networking_faq_sample.txt"
    path.write_text(
        """Networking FAQ

What is DNS?
DNS translates domain names into IP addresses.

What is localhost?
localhost usually refers to 127.0.0.1.

What is ping?
Ping checks connectivity between hosts.

What is a VPN?
A VPN creates a secure encrypted connection.
""",
        encoding="utf-8"
    )


def create_md():
    path = DATA_DIR / "ai_notes_sample.md"
    path.write_text(
        """# AI Notes

RAG means Retrieval Augmented Generation.

Embeddings convert text into numerical vectors.

FAISS is used for similarity search between embeddings.

The AI File Assistant answers questions based on uploaded files.
""",
        encoding="utf-8"
    )


def create_json():
    path = DATA_DIR / "project_info_sample.json"
    data = {
        "project": "AI File Assistant",
        "goal": "Answer questions based on multiple file types",
        "supported_file_types": [
            "txt",
            "md",
            "json",
            "csv",
            "pdf",
            "docx",
            "xlsx"
        ],
        "technologies": [
            "Flask",
            "RAG",
            "FAISS",
            "Gemini",
            "HuggingFace"
        ]
    }

    path.write_text(
        json.dumps(data, indent=2),
        encoding="utf-8"
    )


def create_csv():
    path = DATA_DIR / "students_sample.csv"

    df = pd.DataFrame([
        {"name": "Amit", "course": "AI Engineering", "grade": 92, "status": "passed"},
        {"name": "Yovel", "course": "AI Engineering", "grade": 76, "status": "passed"},
        {"name": "Dana", "course": "Python", "grade": 58, "status": "failed"},
        {"name": "Noa", "course": "Docker", "grade": 84, "status": "passed"},
        {"name": "Eli", "course": "Flask", "grade": 67, "status": "passed"},
    ])

    df.to_csv(path, index=False)


def create_xlsx():
    path = DATA_DIR / "sample_inventory.xlsx"

    df = pd.DataFrame([
        {
            "sku": "SKU-1001",
            "product_name": "USB-C Cable",
            "category": "Accessories",
            "quantity": 120,
            "price_usd": 6.99,
            "status": "in_stock"
        },
        {
            "sku": "SKU-1002",
            "product_name": "Wireless Mouse",
            "category": "Peripherals",
            "quantity": 45,
            "price_usd": 18.50,
            "status": "in_stock"
        },
        {
            "sku": "SKU-1003",
            "product_name": "Mechanical Keyboard",
            "category": "Peripherals",
            "quantity": 18,
            "price_usd": 74.90,
            "status": "low_stock"
        },
        {
            "sku": "SKU-1004",
            "product_name": "HDMI Adapter",
            "category": "Accessories",
            "quantity": 0,
            "price_usd": 12.00,
            "status": "out_of_stock"
        },
    ])

    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="Inventory", index=False)


def create_docx():
    path = DATA_DIR / "sample_policy.docx"

    document = Document()

    document.add_heading("AI File Assistant Policy", 0)

    document.add_paragraph(
        "Purpose: The system should answer questions based on uploaded documents."
    )

    document.add_heading("Supported Document Types", level=1)

    items = [
        "PDF files are read with pypdf.",
        "DOCX files are read with python-docx.",
        "Excel files are read with pandas and openpyxl.",
        "CSV and JSON files are converted into readable text.",
    ]

    for item in items:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Rules", level=1)
    document.add_paragraph(
        "The assistant should not invent facts that are not found in the documents."
    )
    document.add_paragraph(
        "The user interface should show retrieved context to prove that RAG is being used."
    )

    document.save(path)


def create_pdf():
    path = DATA_DIR / "sample_project_guide.pdf"

    pdf = canvas.Canvas(str(path), pagesize=A4)

    width, height = A4
    x = 50
    y = height - 50

    lines = [
        "AI File Assistant - Project Guide",
        "",
        "This PDF is a test document for the RAG ingestion pipeline.",
        "",
        "Main goal:",
        "Answer questions based on uploaded files using Flask, RAG, FAISS, Gemini, and HuggingFace.",
        "",
        "PDF support:",
        "PDF files are extracted with the pypdf library.",
        "",
        "DOCX support:",
        "DOCX files are extracted with the python-docx library.",
        "",
        "Excel support:",
        "XLSX files are extracted with pandas and openpyxl.",
        "",
        "Demo question:",
        "What library is used to read PDF files?",
    ]

    pdf.setFont("Helvetica", 12)

    for line in lines:
        pdf.drawString(x, y, line)
        y -= 20

    pdf.save()


def main():
    create_txt()
    create_md()
    create_json()
    create_csv()
    create_xlsx()
    create_docx()
    create_pdf()

    print("Sample files created successfully inside the data folder.")


if __name__ == "__main__":
    main()